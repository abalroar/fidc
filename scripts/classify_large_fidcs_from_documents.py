"""Read CVM/FundosNet documents for every FIDC above a PL threshold.

The output is deliberately auditable: one row per listed document and one row
per large fund. Raw documents and extracted text stay in the ignored cache;
only the classification ledger is versioned.
"""

from __future__ import annotations

import argparse
from collections import Counter
from datetime import datetime, timezone
from io import BytesIO
import json
from pathlib import Path
import re
import sys
import logging

import pandas as pd
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.classify_fidc_sectors_and_practices import RULES, classify_texts, normalize_text  # noqa: E402
from services.fundonet_client import FundosNetClient  # noqa: E402
from services.regulatory_knowledge import classify_document, should_download_document  # noqa: E402


DEFAULT_THRESHOLD = 5_000_000_000.0
DOC_PRIORITY = {"regulamento": 5, "emissao": 4, "assembleia": 3, "evento": 2, "outro": 1}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--industry-dir", type=Path, default=Path("data/industry_study"))
    parser.add_argument("--raw-dir", type=Path, default=Path("data/raw/industry_large_funds"))
    parser.add_argument("--threshold", type=float, default=DEFAULT_THRESHOLD)
    parser.add_argument("--timeout", type=int, default=60)
    parser.add_argument("--max-docs-per-fund", type=int, default=0, help="0 lê todos os documentos classificatórios")
    parser.add_argument("--force", action="store_true")
    return parser.parse_args()


def safe_name(value: object) -> str:
    text = re.sub(r"[^A-Za-z0-9._-]+", "_", str(value or "").strip()).strip("._")
    return text[:120] or "documento"


def detect_format(content: bytes, declared_name: str) -> str:
    if content.startswith(b"%PDF"):
        return "pdf"
    if content.startswith(b"PK\x03\x04"):
        return "docx"
    suffix = Path(declared_name).suffix.lower().lstrip(".")
    return suffix if suffix in {"pdf", "docx", "doc", "txt", "html", "htm"} else "bin"


def extract_pdf(content: bytes) -> tuple[str, int]:
    reader = PdfReader(BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            parts.append("")
    return "\n".join(parts), len(reader.pages)


def extract_docx(content: bytes) -> tuple[str, int]:
    from docx import Document

    document = Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts), 0


def extract_text(content: bytes, file_format: str) -> tuple[str, int, str]:
    if file_format == "pdf":
        text, pages = extract_pdf(content)
        return text, pages, "pypdf"
    if file_format == "docx":
        text, pages = extract_docx(content)
        return text, pages, "python-docx"
    if file_format in {"txt", "html", "htm"}:
        return content.decode("utf-8", errors="replace"), 0, "texto"
    return "", 0, "formato_não_suportado"


def evidence_snippet(text: str, rule_id: str) -> str:
    normalized = normalize_text(text)
    rule = next((item for item in RULES if item.rule_id == rule_id), None)
    if rule is None:
        return ""
    for pattern in rule.patterns:
        match = re.search(pattern, normalized)
        if match:
            start = max(0, match.start() - 120)
            end = min(len(normalized), match.end() + 180)
            return normalized[start:end]
    return ""


def first_context(text: str, patterns: list[str], *, before: int = 100, after: int = 420) -> tuple[str, str]:
    normalized = normalize_text(text)
    for rule_id, pattern in enumerate(patterns, start=1):
        match = re.search(pattern, normalized)
        if match:
            start = max(0, match.start() - before)
            end = min(len(normalized), match.end() + after)
            return f"large_fidc_{rule_id}", normalized[start:end]
    return "", ""


def extract_anbima_labels(text: str) -> tuple[str, str, str]:
    normalized = normalize_text(text)
    match = re.search(
        r"(?:TIPO ANBIMA|TIPO)\s+[\"']?([A-Z ,/]+?)[\"']?\s+(?:,?\s*COM\s+)?FOCO DE ATUACAO\s+(?:EM\s+)?[\"']?([A-Z ,/]+?)[\"']?(?:\s*,|\s+NOS TERMOS|\s+CONFORME)",
        normalized,
    )
    if not match:
        return "", "", ""
    return match.group(1).strip(), match.group(2).strip(), match.group(0)[:600]


def classify_large_fund_precise(name: str, text: str, ime_segment: str) -> dict[str, str]:
    joined = f"{name}\n{text}"
    rules: list[tuple[str, str, str, list[str]]] = [
        (
            "Meios de Pagamento e Cartões",
            "Arranjos de pagamento/adquirência",
            "meios_pagamento_documental",
            [
                r"TRANSACOES DE PAGAMENTO REALIZADAS NO AMBITO DOS ARRANJOS DE PAGAMENTO",
                r"CLOUDWALK|PAGSEGURO|\bCIELO\b|\bTAPSO\b|\bMONEE\b",
            ],
        ),
        (
            "Crédito PF",
            "Consignado/INSS",
            "consignado_documental",
            [r"CONSIGNACAO EM FOLHA DE PAGAMENTO|CONSIGNADO PRIVADO|CONVENIO INSS|\bCONSIGNADOS?\b"],
        ),
        (
            "Crédito PF",
            "Auto/Veículos",
            "auto_documental",
            [r"PAN AUTO|FINANCIAMENTO DE VEICULOS|ALIENACAO FIDUCIARIA DE VEICUL"],
        ),
        (
            "Judicial/Precatórios/NPL",
            "Precatórios/direitos judiciais",
            "judicial_documental",
            [
                r"DIREITOS CREDITORIOS SAO ORIUNDOS DE LITIGIOS",
                r"ORIGINARIOS DE PRECATORIOS|RESULTEM DE ACOES JUDICIAIS",
                r"RIO VERMELHO.*NAO PADRONIZADO|ALTERNATIVE ASSETS III",
            ],
        ),
        (
            "Infra/Energia",
            "Energia/infra",
            "energia_documental",
            [r"XP COMERCIALIZADORA DE ENERGIA|CAMARA DE COMERCIALIZACAO DE ENERGIA|\bAETOS ENERGIA\b"],
        ),
        (
            "Crédito PJ",
            "Recebíveis comerciais/multissetorial",
            "petrobras_recebiveis_documental",
            [r"DIREITOS CREDITORIOS ORIGINARIOS DE OPERACOES REALIZADAS POR EMPRESAS DO SISTEMA PETROBRAS"],
        ),
        (
            "Crédito PJ",
            "Crédito privado/mercado de capitais",
            "credito_privado_documental",
            [r"ITAU CREDITO PRIVADO", r"DEBENTURE.*CERTIFICADOS DE RECEBIVEIS IMOBILIARIOS.*CERTIFICADOS DE RECEBIVEIS DO AGRONEGOCIO"],
        ),
        (
            "Multissetorial / Outros",
            "Multicarteira outros",
            "multicarteira_documental",
            [
                r"FOCO DE ATUACAO EM MULTICARTEIRA(?:S)? OUTROS",
                r"SEM LIMITACAO QUANTO AOS SEGMENTOS ECONOMICOS",
                r"QUAISQUER DIREITOS CREDITORIOS, INCLUSIVE AQUELES CONSIDERADOS COMO NAO PADRONIZADOS",
            ],
        ),
        (
            "Agro",
            "Agro",
            "agro_documental",
            [r"CEDULA DE PRODUTO RURAL|DIREITOS CREDITORIOS.*PRODUTOR(?:ES)? RURAIS|CADEIAS PRODUTIVAS DO AGRONEGOCIO"],
        ),
    ]
    for n1, n2, rule, patterns in rules:
        _context_rule, context = first_context(joined, patterns)
        if context:
            return {"n1": n1, "n2": n2, "rule": rule, "confidence": "alta", "evidence": context}

    ime_map = {
        "Industrial": ("Crédito PJ", "Recebíveis comerciais/multissetorial"),
        "Comercial": ("Crédito PJ", "Recebíveis comerciais/multissetorial"),
        "Servicos": ("Crédito PJ", "Recebíveis comerciais/multissetorial"),
        "Agronegocio": ("Agro", "Agro"),
        "Acoes judiciais": ("Judicial/Precatórios/NPL", "Precatórios/direitos judiciais"),
        "Setor publico": ("Judicial/Precatórios/NPL", "Precatórios/direitos judiciais"),
        "Cartao de credito": ("Meios de Pagamento e Cartões", "Arranjos de pagamento/adquirência"),
    }
    if ime_segment in ime_map:
        n1, n2 = ime_map[ime_segment]
        return {
            "n1": n1,
            "n2": n2,
            "rule": "segmento_ime_cvm",
            "confidence": "média",
            "evidence": f"Tabela II do Informe Mensal CVM: {ime_segment}",
        }
    return {
        "n1": "Não classificado",
        "n2": "Revisar manualmente",
        "rule": "sem_evidencia_especifica",
        "confidence": "baixa",
        "evidence": "Regulamentos lidos sem política setorial suficientemente específica.",
    }


def classify_fund(name: str, document_texts: list[tuple[str, str]]) -> tuple[object, dict[str, int], str]:
    votes: Counter[tuple[str, str]] = Counter()
    evidence: list[str] = []
    for classification, text in document_texts:
        result = classify_texts({"nome_fundo": name, "observacao_tecnica": text})
        if result.rule_id == "unclassified":
            continue
        weight = DOC_PRIORITY.get(classification, 1)
        votes[(result.setor_n1, result.setor_n2)] += weight
        snippet = evidence_snippet(text, result.rule_id)
        if snippet:
            evidence.append(f"{classification}: {snippet}")
    combined = "\n".join(text for _classification, text in document_texts)
    result = classify_texts({"nome_fundo": name, "observacao_tecnica": combined})
    if votes:
        winner = votes.most_common(1)[0][0]
        if (result.setor_n1, result.setor_n2) != winner:
            winner_text = " ".join(text for classification, text in document_texts if classification in DOC_PRIORITY)
            winner_result = classify_texts({"nome_fundo": name, "observacao_tecnica": winner_text})
            if (winner_result.setor_n1, winner_result.setor_n2) == winner:
                result = winner_result
    return result, {f"{key[0]} | {key[1]}": value for key, value in votes.items()}, " | ".join(evidence[:4])[:1800]


def main() -> None:
    args = parse_args()
    logging.getLogger("pypdf").setLevel(logging.ERROR)
    universe = pd.read_csv(args.industry_dir / "universe_latest.csv", dtype={"cnpj": str}, low_memory=False)
    universe["pl"] = pd.to_numeric(universe["pl"], errors="coerce").fillna(0)
    large = universe[universe["pl"].ge(args.threshold)].sort_values("pl", ascending=False).copy()
    args.raw_dir.mkdir(parents=True, exist_ok=True)
    client = FundosNetClient(timeout_seconds=args.timeout, max_retries=2)

    document_rows: list[dict[str, object]] = []
    fund_rows: list[dict[str, object]] = []
    for position, fund in enumerate(large.itertuples(index=False), start=1):
        cnpj = re.sub(r"\D", "", str(fund.cnpj)).zfill(14)
        name = str(fund.denominacao)
        print(f"[{position}/{len(large)}] {name[:78]}", flush=True)
        try:
            documents = client.listar_documentos(cnpj, error_stage="listar_documentos_large_fidc")
            listing_error = ""
        except Exception as exc:  # noqa: BLE001
            documents = []
            listing_error = f"{type(exc).__name__}: {exc}"

        selected = []
        for document in documents:
            classification = classify_document(
                categoria=document.categoria,
                tipo=document.tipo,
                especie=document.especie,
                nome_arquivo=document.nome_arquivo or "",
            )
            if should_download_document(classification):
                selected.append((document, classification))
        selected.sort(
            key=lambda item: (
                DOC_PRIORITY.get(item[1], 0),
                item[0].data_referencia_dt or datetime.min.date(),
                item[0].versao,
                item[0].id,
            ),
            reverse=True,
        )
        if args.max_docs_per_fund > 0:
            selected = selected[: args.max_docs_per_fund]
        selected_ids = {document.id for document, _classification in selected}

        fund_dir = args.raw_dir / cnpj
        fund_dir.mkdir(parents=True, exist_ok=True)
        extracted: list[tuple[str, str]] = []
        read_docs = 0
        text_docs = 0
        total_chars = 0
        for document, classification in selected:
            declared_name = document.nome_arquivo or f"{classification}_{document.id}.pdf"
            base_path = fund_dir / f"{document.id}_{classification}_{safe_name(declared_name)}"
            content = b""
            download_error = ""
            try:
                existing = next(fund_dir.glob(f"{document.id}_{classification}_*"), None)
                if existing is not None and existing.stat().st_size > 0 and not args.force:
                    content = existing.read_bytes()
                    base_path = existing
                else:
                    content = client.download_documento(document.id)
                    file_format = detect_format(content, declared_name)
                    base_path = base_path.with_suffix(f".{file_format}")
                    base_path.write_bytes(content)
            except Exception as exc:  # noqa: BLE001
                download_error = f"{type(exc).__name__}: {exc}"

            file_format = detect_format(content, declared_name) if content else ""
            text = ""
            pages = 0
            extraction_method = ""
            extraction_error = ""
            if content:
                try:
                    text, pages, extraction_method = extract_text(content, file_format)
                    read_docs += 1
                except Exception as exc:  # noqa: BLE001
                    extraction_error = f"{type(exc).__name__}: {exc}"
            normalized_text = re.sub(r"\s+", " ", text).strip()
            if normalized_text:
                text_docs += 1
                total_chars += len(normalized_text)
                extracted.append((classification, normalized_text[:1_500_000]))
                text_path = fund_dir / f"{document.id}_{classification}.txt"
                if args.force or not text_path.exists():
                    text_path.write_text(normalized_text, encoding="utf-8")

            document_rows.append(
                {
                    "competencia_snapshot": fund.competencia,
                    "cnpj": cnpj,
                    "fund_name": name,
                    "pl_brl": float(fund.pl),
                    "document_id": document.id,
                    "categoria": document.categoria,
                    "tipo": document.tipo,
                    "especie": document.especie,
                    "document_classification": classification,
                    "selected_for_classification": True,
                    "data_referencia": document.data_referencia or "",
                    "data_entrega": document.data_entrega or "",
                    "versao": document.versao,
                    "status": document.status,
                    "declared_file_name": declared_name,
                    "file_format": file_format,
                    "file_size_bytes": len(content),
                    "pages": pages,
                    "text_chars": len(normalized_text),
                    "read_status": "texto_extraído" if normalized_text else ("lido_sem_texto" if content else "erro_download"),
                    "extraction_method": extraction_method,
                    "download_error": download_error,
                    "extraction_error": extraction_error,
                }
            )

        for document in documents:
            if document.id in selected_ids:
                continue
            classification = classify_document(
                categoria=document.categoria,
                tipo=document.tipo,
                especie=document.especie,
                nome_arquivo=document.nome_arquivo or "",
            )
            document_rows.append(
                {
                    "competencia_snapshot": fund.competencia,
                    "cnpj": cnpj,
                    "fund_name": name,
                    "pl_brl": float(fund.pl),
                    "document_id": document.id,
                    "categoria": document.categoria,
                    "tipo": document.tipo,
                    "especie": document.especie,
                    "document_classification": classification,
                    "selected_for_classification": False,
                    "data_referencia": document.data_referencia or "",
                    "data_entrega": document.data_entrega or "",
                    "versao": document.versao,
                    "status": document.status,
                    "declared_file_name": document.nome_arquivo or "",
                    "file_format": "",
                    "file_size_bytes": 0,
                    "pages": 0,
                    "text_chars": 0,
                    "read_status": "fora_escopo_classificatório",
                    "extraction_method": "",
                    "download_error": "",
                    "extraction_error": "",
                }
            )

        result, votes, evidence = classify_fund(name, extracted)
        combined_text = "\n".join(text for _classification, text in extracted)
        precise = classify_large_fund_precise(name, combined_text, str(getattr(fund, "segmento_principal", "") or ""))
        anbima_type, anbima_focus, anbima_evidence = extract_anbima_labels(combined_text)
        fund_rows.append(
            {
                "competencia_snapshot": fund.competencia,
                "cnpj": cnpj,
                "fund_name": name,
                "pl_brl": float(fund.pl),
                "pl_share_large_funds": float(fund.pl / large["pl"].sum()),
                "ime_segment": getattr(fund, "segmento_principal", ""),
                "anbima_type_document": anbima_type,
                "anbima_focus_document": anbima_focus,
                "anbima_evidence": anbima_evidence,
                "document_segment_n1": precise["n1"],
                "document_segment_n2": precise["n2"],
                "classification_confidence": precise["confidence"],
                "classification_rule": precise["rule"],
                "classification_evidence": precise["evidence"],
                "generic_classifier_result": f"{result.setor_n1} | {result.setor_n2}",
                "generic_classifier_evidence": evidence or result.evidence,
                "classification_votes": json.dumps(votes, ensure_ascii=False, sort_keys=True),
                "documents_listed": len(documents),
                "documents_relevant": len(selected),
                "documents_read": read_docs,
                "documents_with_text": text_docs,
                "text_chars": total_chars,
                "listing_error": listing_error,
                "coverage_status": "completo" if selected and read_docs == len(selected) else ("parcial" if read_docs else "sem_leitura"),
                "source": "CVM/FundosNet; regulamentos, emissões, assembleias e eventos públicos",
            }
        )

    funds = pd.DataFrame(fund_rows).sort_values("pl_brl", ascending=False)
    docs = pd.DataFrame(document_rows).sort_values(["pl_brl", "cnpj", "data_referencia"], ascending=[False, True, False])
    funds.to_csv(args.industry_dir / "industry_large_fund_classification.csv", index=False)
    docs.to_csv(args.industry_dir / "industry_large_fund_documents.csv.gz", index=False, compression="gzip")
    manifest = {
        "schema_version": "large-fidc-document-classification/v1",
        "generated_at_utc": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "threshold_brl": args.threshold,
        "snapshot_competence": str(funds.iloc[0]["competencia_snapshot"]) if not funds.empty else "",
        "funds": int(len(funds)),
        "pl_brl": float(funds["pl_brl"].sum()) if not funds.empty else 0.0,
        "documents_listed": int(funds["documents_listed"].sum()) if not funds.empty else 0,
        "documents_relevant": int(funds["documents_relevant"].sum()) if not funds.empty else 0,
        "documents_read": int(funds["documents_read"].sum()) if not funds.empty else 0,
        "documents_with_text": int(funds["documents_with_text"].sum()) if not funds.empty else 0,
        "funds_high_confidence": int(funds["classification_confidence"].eq("alta").sum()) if not funds.empty else 0,
        "document_ledger_rows": int(len(docs)),
        "scope_note": "Documentos classificatórios: regulamentos, emissões, assembleias e eventos. Informes mensais estruturados não são relidos como PDF.",
    }
    (args.industry_dir / "industry_large_fund_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(
        f"[ok] {len(funds)} fundos; {manifest['documents_read']}/{manifest['documents_relevant']} "
        "documentos classificatórios lidos",
        flush=True,
    )


if __name__ == "__main__":
    main()
